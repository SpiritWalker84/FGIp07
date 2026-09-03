#!/usr/bin/env python3
"""Заполнение рабочей тетради FGIp07 (кейсы 1–3 + сводка + рефлексия)."""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

SRC = Path(__file__).parent / "FGip07_worknote.docx"
OUT = Path(__file__).parent / "FGip07_worknote_filled.docx"


def set_cell(table, row: int, col: int, text: str) -> None:
    table.rows[row].cells[col].text = text


def fill_case(
    doc: Document,
    *,
    order: int,
    summary_row: int,
    meta: int,
    tz: int,
    mvp: int,
    test: int,
    pub: int,
    case: int,
    data: dict,
) -> None:
    s = doc.tables[summary_row]
    set_cell(s, order, 1, data["exchange"])
    set_cell(s, order, 2, data["niche"])
    set_cell(s, order, 3, data["type"])
    set_cell(s, order, 4, data["budget"])
    set_cell(s, order, 5, data["time"])

    m = doc.tables[meta]
    set_cell(m, 0, 1, data["url"])
    set_cell(m, 1, 1, data["client_goal"])
    set_cell(m, 2, 1, data["tracker_type"])
    set_cell(m, 3, 1, data["order_budget"])

    t = doc.tables[tz]
    set_cell(t, 0, 1, data["features"])
    set_cell(t, 1, 1, data["questions"])
    set_cell(t, 2, 1, data["spec"])

    b = doc.tables[mvp]
    set_cell(b, 0, 1, data["project"])
    set_cell(b, 1, 1, data["strategy"])
    set_cell(b, 2, 1, data["prompt"])

    te = doc.tables[test]
    set_cell(te, 0, 1, data["broken"])
    set_cell(te, 1, 1, data["fixes"])
    set_cell(te, 2, 1, data["retest"])

    p = doc.tables[pub]
    set_cell(p, 0, 1, data["style"])
    set_cell(p, 1, 1, data["github"])
    set_cell(p, 2, 1, data["public_url"])

    c = doc.tables[case]
    set_cell(c, 0, 1, data["price"])
    set_cell(c, 1, 1, data["handoff"])
    set_cell(c, 2, 1, data["upsell"])
    set_cell(c, 3, 1, data["fact_time"])


def main() -> None:
    shutil.copy(SRC, OUT)
    doc = Document(OUT)

    set_cell(doc.tables[2], 0, 1, "Ринат")
    set_cell(doc.tables[2], 1, 1, "— (уточнить поток у куратора)")
    set_cell(doc.tables[2], 2, 1, "02.09.2026")
    set_cell(doc.tables[2], 3, 1, "[Вставить ссылку на Яндекс.Диск после загрузки]")

    case1 = {
        "exchange": "FL.ru",
        "niche": "Корпоративный task-трекер для команды",
        "type": "Задачи",
        "budget": "3 000 ₽ (заказ) / 25 000 ₽ (моя оценка MVP)",
        "time": "2,5–3 ч / 2 ч 45 мин",
        "url": "https://www.fl.ru/projects/4242426/sdelat-task-treker-na-php-i-bootstrap.html",
        "client_goal": "Внутренний трекер задач с ролями и вкладками «мои задачи / я постановщик / все задачи», позже — корпоративный портал.",
        "tracker_type": "Задачи (корпоративный)",
        "order_budget": "3 000 ₽",
        "features": "Регистрация и авторизация; роли (user, manager, admin); CRUD задач; исполнитель и постановщик; статусы; комментарии; сроки.",
        "questions": "1) Нужна ли канбан-доска или достаточно списков на MVP? 2) Сколько пользователей одновременно? 3) Нужны ли соисполнители в первой версии? 4) Есть ли брендбук? 5) Где хостинг — свой сервер или облако?",
        "spec": "info/tz-case1-taskhub.md — цель: MVP task-трекера; роли admin/manager/user; экраны: вход, список задач, карточка, создание/редактирование; критерий MVP: авторизация + задачи + назначение + комментарии.",
        "project": "FGIp07/case1-web-tasks/ + github.com/SpiritWalker84/TaskHub (Laravel 11, Docker)",
        "strategy": "[ ] один большой промпт  [✓] поэтапно (сначала каркас TaskHub, затем доработки по ТЗ заказа)",
        "prompt": "Разработай MVP трекера задач по ТЗ в папке info. Роли: admin, manager, user. Задачи: title, description, status (new/in_progress/done), assignee, due_date, comments. Laravel + Docker. Сначала рабочий CRUD и авторизация, потом канбан.",
        "broken": "Первый запуск: не было сидов пользователей; список задач без фильтра «мои / все»; канбан не реализован (только список).",
        "fixes": "1) «Добавь сиды demo-пользователей и 5 задач для теста ролей». 2) «Доработай фильтры списка: мои задачи / где я постановщик / все (admin)». 3) «Добавь простую канбан-доску с колонками по статусам».",
        "retest": "Вход под admin и user; создание и редактирование задачи; назначение исполнителя; комментарий; смена статуса; проверка прав на удаление.",
        "style": "Нейтральный минимализм (без брендбука у заказчика): светлая тема, акцент #2563eb, системный шрифт.",
        "github": "https://github.com/SpiritWalker84/TaskHub",
        "public_url": "http://localhost:8086 (Docker). Демо: admin@taskhub.local / password",
        "price": "25 000 ₽ (MVP без канбана в заказе — 15 000 ₽; с канбаном и деплоем — 25–35 000 ₽)",
        "handoff": "GitHub + инструкция Docker; опционально — поддержка сервера у клиента.",
        "upsell": "Поддержка сервера 2–3 тыс. ₽/мес; доработка канбана и уведомлений; интеграция с Telegram.",
        "fact_time": "2 ч 45 мин   Портфолио: [✓] да  [ ] нет",
    }

    case2 = {
        "exchange": "FL.ru",
        "niche": "Task manager в Telegram для команды",
        "type": "Задачи (Telegram-бот)",
        "budget": "по договорённости (отклики 20–50 тыс. ₽)",
        "time": "2–2,5 ч / 2 ч 10 мин",
        "url": "https://www.fl.ru/projects/5450726/task-manager---crm-v-telegram-bot.html",
        "client_goal": "Автономный Telegram-бот (исходники), MySQL, постановка/контроль задач команды, желательно канбан. Не конструктор.",
        "tracker_type": "Задачи",
        "order_budget": "по договорённости",
        "features": "Создание/просмотр/редактирование задач; уведомления о выполнении; MySQL; исходный код; опционально канбан в боте.",
        "questions": "1) Сколько человек в команде? 2) Нужны ли роли (админ/сотрудник)? 3) Канбан — inline-кнопки или Web App? 4) Нужны ли дедлайны и напоминания? 5) Где будет крутиться бот (VPS/Railway)?",
        "spec": "info/tz-case2-telegram-tasks.md — цель: таск-бот для команды; роли admin/member; команды /tasks, /new, /done; статусы todo/in_progress/done; MySQL.",
        "project": "FGIp07/case2-telegram-tasks/ — Python aiogram 3 + MySQL",
        "strategy": "[ ] один большой промпт  [✓] поэтапно (бот-каркас → CRUD задач → уведомления → канбан-кнопки)",
        "prompt": "Создай Telegram-бот таск-менеджер: aiogram 3, MySQL, таблицы users/tasks. Команды: список задач, создать (FSM), сменить статус, назначить исполнителя. Уведомление постановщику при done. Исходники для автономного запуска.",
        "broken": "После первой генерации: FSM создания задачи не сбрасывался; не было разграничения «чужая задача»; уведомления не отправлялись.",
        "fixes": "1) «Исправь сброс состояния FSM после создания задачи». 2) «Добавь проверку: редактировать может постановщик или admin». 3) «При status=done отправляй сообщение creator_id».",
        "retest": "Два аккаунта Telegram; создание задачи; смена статуса; уведомление; список /tasks; роль admin видит все задачи.",
        "style": "Стандартный интерфейс Telegram, emoji-статусы (🆕 🔄 ✅), без отдельного брендбука.",
        "github": "https://github.com/SpiritWalker84/FGIp07 (case2-telegram-tasks/)",
        "public_url": "Docker: case2-telegram-tasks (нужен BOT_TOKEN от @BotFather). Запуск локально или на сервере клиента.",
        "price": "22 000 ₽ (бот + MySQL + инструкция); с канбан WebApp — 30 000 ₽",
        "handoff": "Исходники на GitHub + .env.example + docker-compose; инструкция запуска.",
        "upsell": "Хостинг бота; доработка Web App канбана; интеграция с Bitrix24.",
        "fact_time": "2 ч 10 мин   Портфолио: [✓] да  [ ] нет",
    }

    case3 = {
        "exchange": "Kwork + свой кейс",
        "niche": "Трекер заявок для сервисной компании (диспетчер / мастер)",
        "type": "Задачи / заявки",
        "budget": "учебный кейс (аналог заказов на CRM/заявки)",
        "time": "2–2,5 ч / 2 ч 20 мин",
        "url": "https://kwork.ru/projects?c=11 (поиск CRM/заявок) + реализация: github.com/SpiritWalker84/Zayavki",
        "client_goal": "Принимать заявки с сайта, назначать мастерам, контролировать статус выезда — без Excel и хаоса в мессенджерах.",
        "tracker_type": "Задачи (заявки в сервисе)",
        "order_budget": "—",
        "features": "Публичная форма заявки; роли dispatcher/master; диспетчерская (таблица + назначение); панель мастера «Мои выезды»; статусы new/assigned/in_progress/done; защита от двойного взятия (race condition); UI бренда «МастерДом».",
        "questions": "1) Какие статусы заявки? 2) Нужны ли SMS клиенту? 3) Один или несколько мастеров на заявку? 4) Нужен ли мобильный вид для мастера на объекте? 5) Интеграция с 1С?",
        "spec": "info/tz-case3-zayavki.md + Zayavki README — экраны: /requests/create, /dispatcher, /master; MVP: new → assigned → in_progress → done.",
        "project": "FGIp07/case3-service-requests/ (форк Zayavki, Laravel 11, MySQL, Docker, порт 8087)",
        "strategy": "[✓] поэтапно (Docker локально → редизайн UI «МастерДом» → правки UX профиля/карточек → тест ролей и скриншоты)",
        "prompt": "Доработай трекер заявок под бренд «МастерДом»: split-screen логин, sidebar для ролей, карточки выездов у мастера (адрес и телефон в приоритете), диспетчерская-таблица. Сохрани takeRequest с транзакцией против race condition.",
        "broken": "UI: в профиле дублировалась роль («Мастер Петров» + chip «Мастер»), аватар показывал одну букву «М»; сайдбар на мобилке ломался; у мастера было лишнее меню «Новая заявка». Карточки заявок были перегружены секциями.",
        "fixes": "1) «Раздели displayName и roleLabel, инициалы из ФИО (ПП)». 2) «Упрости карточки выездов: адрес заголовком, клиент/телефон в одном блоке». 3) «Навигация по роли: мастер — только Мои выезды». 4) «Убери декоративную палитру с экрана логина».",
        "retest": "Создание заявки без входа; вход dispatcher — назначение мастера; вход master1 — взять в работу и завершить; две вкладки на одну заявку (одна ошибка); скриншоты: login, create, dispatcher, master.",
        "style": "Бренд «МастерДом»: бирюза #0f766e + янтарный акцент, шрифт Plus Jakarta Sans, тёмный sidebar, split-screen авторизация, компактные карточки выездов (отличие от синего TaskHub в кейсе №1).",
        "github": "https://github.com/SpiritWalker84/Zayavki (база) + FGIp07/case3-service-requests/",
        "public_url": "http://localhost:8087 (Docker). Демо: dispatcher@example.com / master1@example.com — пароль password",
        "price": "22 000 ₽ (MVP заявок + деплой + UI под бренд; с SMS/Telegram — 30 000 ₽)",
        "handoff": "GitHub + docker-compose; демо-логины; краткая инструкция для диспетчера.",
        "upsell": "SMS/Telegram уведомления клиенту; отчёты по мастерам; PWA для мастера на объекте.",
        "fact_time": "2 ч 20 мин   Портфолио: [ ] да  [✓] нет (в портфолио — кейс №1)",
    }

    fill_case(doc, order=1, summary_row=5, meta=9, tz=11, mvp=13, test=15, pub=17, case=20, data=case1)
    fill_case(doc, order=2, summary_row=5, meta=23, tz=25, mvp=27, test=29, pub=31, case=34, data=case2)
    fill_case(doc, order=3, summary_row=5, meta=37, tz=39, mvp=41, test=43, pub=45, case=48, data=case3)

    prog = doc.tables[149]
    set_cell(prog, 1, 1, "2 ч 45 мин")
    set_cell(prog, 1, 2, "Уже был каркас TaskHub — быстрее ТЗ и промпты")
    set_cell(prog, 2, 1, "2 ч 10 мин")
    set_cell(prog, 2, 2, "Шаблон aiogram из telegram-bots-portfolio")
    set_cell(prog, 3, 1, "2 ч 20 мин")
    set_cell(prog, 3, 2, "Каркас Zayavki + Docker + редизайн UI «МастерДом»")

    set_cell(
        doc.tables[150],
        1,
        0,
        "Анализ ТЗ и вопросы клиенту — быстрее: после трёх кейсов сразу выделяю роли, статусы и границу MVP. "
        "Промпты — поэтапные вместо одного большого; так быстрее получаю рабочий код. "
        "Деплой через Docker (migrate, seed, порты) — уже по шаблону, без долгого разбора. "
        "Тестирование — автоматически проверяю 2 роли, CRUD, смену статуса и граничные случаи (race condition в кейсе №3). "
        "Публикация — README, демо-логины и URL в одном формате.",
    )
    set_cell(
        doc.tables[151],
        1,
        0,
        "Во всех трёх кейсах повторялись: авторизация, роли, статусы, назначение исполнителя, списки/карточки. "
        "В заявках и задачах — уведомления и контроль «кто взял в работу». "
        "Реже в бюджетных заказах: канбан, AI, интеграции с 1С — это уже допродажа.",
    )
    set_cell(
        doc.tables[152],
        1,
        0,
        "Кейс №1 TaskHub — самый сильный для портфолио: полный веб-трекер под FL.ru, канбан, фильтры, Docker. "
        "№2 (Telegram) — хорош как второй формат (бот + MySQL). №3 (МастерДом) — учебный, но показывает заявки и UI под бренд.",
    )

    set_cell(doc.tables[156], 0, 1, "[Яндекс.Диск — после загрузки тетради и скриншотов]")
    set_cell(doc.tables[156], 1, 1, "3")
    set_cell(doc.tables[156], 2, 1, "№1 и №2")
    set_cell(doc.tables[156], 3, 1, "02.09.2026")

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
